"""
Builds data/sample_documents/ - real PDF and DOCX files for exercising the
document-input path in the app.

Run:  python scripts/fetch_sample_documents.py

These are separate from data/test_documents/, which holds the *annotated*
corpus used for scoring. These are unannotated, real-world files whose only
job is to prove that PDF and DOCX extraction works on documents nobody
prepared for us.

PDFs   real SEC EDGAR contracts from CUAD (CC BY 4.0), used as filed
DOCX   generated from our own synthetic HR thread, because there is no
       public corpus of real HR grievance documents - and there should not
       be. Written with a table, since DOCX tables are the input path most
       likely to break silently.
"""

from __future__ import annotations

from pathlib import Path

from huggingface_hub import hf_hub_download

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "data" / "sample_documents"

# Chosen for variety of contract type and length, not cherry-picked for results.
CUAD_PDFS = {
    "service_agreement.pdf":
        "CUAD_v1/full_contract_pdf/Part_I/Service/"
        "AMERIQUESTTECHNOLOGIESINC_04_27_2018-EX-10.1-SERVICE AGREEMENT.pdf",
    "maintenance_agreement.pdf":
        "CUAD_v1/full_contract_pdf/Part_I/Maintenance/"
        "CREDITCARDSCOMINC_08_10_2007-EX-10.33-Maintenance Agreement.pdf",
    "affiliate_agreement.pdf":
        "CUAD_v1/full_contract_pdf/Part_I/Affiliate_Agreements/"
        "LinkPlusCorp_20050802_8-K_EX-10_3240252_EX-10_Affiliate Agreement.pdf",
}


BASE_URL = "https://huggingface.co/datasets/theatticusproject/cuad/resolve/main/"


def fetch_pdfs() -> list[Path]:
    """Download real contract PDFs straight to short local filenames.

    Deliberately NOT via hf_hub_download: CUAD's filenames are long, and the
    cache nests them under a snapshot hash, which pushes the full path past
    Windows' 260-character MAX_PATH limit and fails with a bare
    FileNotFoundError. Streaming to our own short name sidesteps that.
    """
    import requests
    from urllib.parse import quote
    from huggingface_hub import list_repo_files

    available = [f for f in list_repo_files("theatticusproject/cuad",
                                            repo_type="dataset")
                 if f.endswith(".pdf") and "full_contract_pdf" in f]
    written: list[Path] = []

    for name, wanted in CUAD_PDFS.items():
        source = wanted if wanted in available else None
        if source is None:
            # Match on contract type so we still get a comparable document.
            category = wanted.split("/")[3]
            candidates = [f for f in available if f"/{category}/" in f]
            source = candidates[0] if candidates else available[0]
            print(f"  ! {name}: listed path unavailable, using "
                  f"{source.split('/')[-1][:46]}")

        target = OUT / name
        if target.exists() and target.stat().st_size > 10_000:
            print(f"  = {name:28s} already present, skipping")
            written.append(target)
            continue

        # Unauthenticated Hub requests are rate limited; 429 is routine when
        # pulling several files in a row, so back off rather than give up.
        import time
        for attempt in range(5):
            response = requests.get(BASE_URL + quote(source), timeout=120)
            if response.status_code != 429:
                break
            wait = 5 * (attempt + 1)
            print(f"    rate limited, retrying in {wait}s ...")
            time.sleep(wait)
        response.raise_for_status()
        target.write_bytes(response.content)
        written.append(target)
        print(f"  + {name:28s} {target.stat().st_size/1024:7.1f} KB   "
              f"<- {source.split('/')[-1][:44]}")
    return written


def build_docx() -> Path:
    """Write a real .docx containing paragraphs AND a table.

    The table matters: contract schedules and HR forms keep much of their
    sensitive detail in tables, and a paragraph-only reader drops it without
    complaint. This file exercises that path.
    """
    from docx import Document

    doc = Document()
    doc.add_heading("Employee Grievance Record", level=1)
    doc.add_paragraph(
        "This record summarises a formal grievance raised under the company "
        "grievance procedure. It is retained by People Operations and is "
        "subject to applicable data-protection law.")

    doc.add_heading("Complainant details", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Value"
    for field, value in [
        ("Name", "Aisha Rahman"),
        ("Employee ID", "EC-2019-0884"),
        ("Work email", "aisha.rahman@example-corp.com"),
        ("Personal email", "aisha.r.personal@example.net"),
        ("Mobile", "+31 6 5550 2277"),
        ("Home address", "48 Weena Zuid, Rotterdam 3012 NC"),
        ("Line manager", "Thomas Whitfield"),
        ("Case reference", "GRV-2026-0417"),
        ("HR contact", "Neha Okonkwo"),
    ]:
        row = table.add_row()
        row.cells[0].text = field
        row.cells[1].text = value

    doc.add_heading("Summary of complaint", level=2)
    doc.add_paragraph(
        "The complainant reports that comments made in team meetings were "
        "inappropriate. The matter was raised informally and the complainant "
        "was told it would be handled appropriately, but states that nothing "
        "much has changed since then.")
    doc.add_paragraph(
        "People Operations will review the matter and aim to respond within a "
        "reasonable timeframe. Where practicable, an initial response will be "
        "provided promptly.")

    doc.add_heading("Retention", level=2)
    doc.add_paragraph(
        "This record shall be retained for six years, and may be destroyed "
        "sooner where no longer necessary. Records must not be destroyed "
        "before the end of the retention period.")

    target = OUT / "hr_grievance_record.docx"
    doc.save(str(target))
    print(f"  + {target.name:28s} {target.stat().st_size/1024:7.1f} KB")
    return target


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    print("Fetching real contract PDFs from CUAD ...")
    pdfs = fetch_pdfs()
    print("\nBuilding DOCX with a table ...")
    docx = build_docx()

    print("\nVerifying extraction ...")

    # app.py is a Streamlit script: importing it executes the UI at module
    # level and crashes outside a Streamlit runtime. So we exercise the same
    # two libraries the app uses (PyMuPDF, python-docx) rather than importing
    # from it.
    def extract_text(name: str, data: bytes) -> str:
        import re as _re
        suffix = name.rsplit(".", 1)[-1].lower()
        if suffix == "pdf":
            import fitz
            with fitz.open(stream=data, filetype="pdf") as d:
                return "\n".join(page.get_text() for page in d)
        if suffix == "docx":
            from io import BytesIO
            from docx import Document
            d = Document(BytesIO(data))
            parts = [p.text.strip() for p in d.paragraphs if p.text.strip()]
            for n, table in enumerate(d.tables, start=1):
                for row in table.rows:
                    cells = list(dict.fromkeys(
                        _re.sub(r"\s+", " ", c.text).strip()
                        for c in row.cells if c.text.strip()))
                    if cells:
                        parts.append(f"[Table {n}] " + " | ".join(cells))
            return "\n".join(parts)
        return data.decode("utf-8", errors="replace")

    ok = True
    for path in pdfs + [docx]:
        try:
            text = extract_text(path.name, path.read_bytes())
            words = len(text.split())
            status = "OK " if words > 50 else "THIN"
            if words <= 50:
                ok = False
            print(f"  {status} {path.name:28s} {words:6,d} words   "
                  f"{text.strip()[:60]!r}")
        except Exception as exc:
            ok = False
            print(f"  FAIL {path.name:28s} {type(exc).__name__}: {exc}")

    print("\nEXTRACTION:", "PASS" if ok else "FAIL")
    if not ok:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
